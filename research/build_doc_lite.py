#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_doc_lite.py — 在 增订版_配图.docx 基础上的「小红书长文版」：
保留 7 幅图 + 3 张表图 + 口语化开头/结论速览/标题层级，但把正文文字压到
1 万字以内（删整段参考文献 -> 一行精选；精简最长的几节）。

文字按块内联在 CONTENT 中，图/表图按块插入，便于精确控字数。
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGDIR = "/sessions/busy-optimistic-gates/mnt/WebstormProjects/paper_figures/figures"
OUT = "/sessions/busy-optimistic-gates/mnt/outputs/LLM多智能体投资系统_增订版_配图.docx"
LATIN, CJK = "Times New Roman", "PingFang SC"
GRAY = RGBColor(0x59, 0x5F, 0x66); DARK = RGBColor(0x20, 0x24, 0x28)


def set_run(run, size=10.5, bold=False, color=None, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = LATIN
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), LATIN); rf.set(qn("w:hAnsi"), LATIN); rf.set(qn("w:eastAsia"), CJK)
    return run


def para(doc, text="", size=10.5, bold=False, align=None, color=None, italic=False,
         space_after=6, space_before=0, line=None, indent_first=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    if line is not None:
        pf.line_spacing = line
    if indent_first is not None:
        pf.first_line_indent = Pt(indent_first)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_figure(doc, fname, cap, width=6.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(FIGDIR, fname), width=Inches(width))
    para(doc, cap, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_before=3, space_after=10)


def add_image(doc, fname, width=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(os.path.join(FIGDIR, fname), width=Inches(width))


# ============================== CONTENT ==============================
INTRO = [
    "先说这是什么：OpenInvest 是我自己捣鼓的一个项目——让几个大语言模型分别当宏观、量化、风控、CIO，"
    "组成一个“投资委员会”，对资产未来 30/60/90 天的涨跌给出带概率的判断。做的过程中我一直纠结：这种多"
    "智能体系统，到底要几个角色、怎么分工才最好？是不是角色越多就越聪明？",
    "于是我做了两件事：把 2024–2026 年这方面的论文翻了一遍，再在自己系统上跑了一批实验来验证。下面是结论。"
    "每个数字我都尽量用真实回测撑着（全历史 1579 个零前视重算点，所有图表都能用脚本复现），但我确实是新手，"
    "思路和判断难免有错——看出哪里不对，特别欢迎在评论区告诉我，让我有机会把这个系统改得更好。",
]
TLDR = [
    ("没有“最优角色数”这种东西。",
     "七角色的 TradingAgents 和我的四角色，在证据上是平级的——没有任何已发表框架真正做过“角色数量”的消融"
     "实验，角色怎么分基本都是“模仿真实公司分工”拍脑袋定的（见表 1）。"),
    ("多智能体的好处主要来自“多采样＋聚合”，不是“角色扮演”本身。",
     "算力对齐的对照实验里，单模型多跑几次再确定性聚合，就能拿到大部分所谓“多智能体”的增益（图 1 机制 b）。"),
    ("让 LLM 自己报概率/置信度不靠谱。",
     "我系统里这些自报置信度几乎全堆在 0.55/0.60/0.65 上（图 7），是已知的“整数化伪精度”，得在 LLM 外面"
     "再加一层确定性校准。"),
    ("辩论不一定有用，还可能更糟。",
     "理论上多轮辩论的期望正确率不变（鞅性质），实际还容易把大家带向“过度自信”，恰好被校准指标扣分。"),
    ("我加的校准层是有效的。",
     "真实数据验证后，30/60/90 天概率带覆盖率回到了 [75%, 85%] 目标区间，Brier 全线改善（图 6、表 2）。"),
    ("能不能赚钱，短期根本测不出来。",
     "统计上夏普 0.5 的策略约要 16 年数据才能确认（图 2）；当下能验证的只有“校准”和“纪律”，不是收益。"),
]

# 每个块: ("h1"|"body"|"fig"|"img", payload)
BLOCKS = [
    ("h1", "1　研究的问题"),
    ("body", "多智能体 LLM 金融系统这两年很火。TradingAgents 用七类角色（基本面/情绪/新闻/技术四类分析师 + 多空"
             "研究员 + 交易员 + 风控团队）模仿真实交易公司；我的 OpenInvest 用四角色（宏观/量化/风控/CIO）、单轮"
             "汇总、不辩论，并在 LLM 之外加了确定性后处理与独立防御组件。两者共享一个没被检验的前提：角色划分"
             "本身有效——前者七角色来自对公司部门的类比，后者四角色来自我的经验直觉，都没经过受控实验筛选。"),
    ("body", "我想回答四个问题：RQ1 角色数量/划分有没有可被实证支持的最优解？RQ2 若没有，该按什么原则设计？"
             "RQ3 校准优先下我系统各组件的证据地位如何、该怎么改？RQ4 校准目标和盈利目标怎么界定，既不混淆也"
             "不割裂？"),
    ("body", "口径很关键：我的主目标是“预测校准质量”（系统报出的概率本身可不可信，用 Brier 和概率带覆盖率衡量），"
             "回测收益是次要目标。两类口径本质不同：收益类评估整条决策链的最终结果，校准类只看“系统所述概率”和"
             "“实际频率”一不一致。这跟 TradingAgents、FinCon 等不同——它们都只报收益类指标，不报任何校准指标。"),

    ("h1", "2　别人做过什么（证据综述）"),
    ("body", "主流框架的实证基础都不牢。TradingAgents 的回测窗口只有约三个月（2024 年 1–3 月）、标的只有 AAPL/"
             "GOOGL/AMZN 三只、对照仅五个规则策略，Sharpe 报到 8.21/6.39/5.60（作者自己在脚注里都说超出合理范围），"
             "且全文不含任何角色数量的消融、不含单 LLM 基线、不含校准指标。FinCon、AlphaAgents 等同样靠“模拟真实"
             "机构分工”来论证角色配置。各框架实证基础对比见表 1。"),
    ("img", ("table1_frameworks.png", 6.5)),
    ("body", "等算力对照下，多智能体并没占到便宜：Wang 等（ACL 2024）发现配强结构化提示的单智能体可匹配最优多"
             "智能体讨论；Smit 等（ICML 2024）发现多智能体辩论不能可靠优于 self-consistency、集成等更简单的策略；"
             "Zhang 等（2025）在九个基准、四个模型上复核，结论是辩论即使多耗算力也无法可靠超过单智能体；Tran 与 "
             "Kiela（2026）从数据处理不等式论证固定预算下单智能体信息效率更高。小结：受控条件下，归因于“角色扮演”"
             "的增益，主要还原成“多次采样＋聚合”的集成效应，不依赖角色划分。"),
    ("body", "辩论的麻烦：Choi 等（NeurIPS 2025）证明辩论让信念轨迹构成鞅——设第 t 轮信念为 B_t，则 "
             "E[B_{t+1}|F_t]=B_t，期望正确率不随辩论轮次提升；只有往纠正方向施加偏置才能打破中性。同时多项工作"
             "记录了辩论的谄媚、从众、过早收敛、被论断性而非证据说服等失效模式，方向都是把群体推向“自信的共识”。"
             "而 Brier 和覆盖率恰好惩罚过度自信，所以辩论和校准优先目标方向冲突。"),
    ("body", "言语化概率不可靠：Xiong 等（ICLR 2024）发现 LLM 自报置信度高度聚在 80–100% 且呈 5 的倍数，平均"
             "校准误差超过 0.377，GPT-4 区分自身对错的 AUROC 只有 62.7%（接近随机）；给模型设“自信/谨慎”人设几乎"
             "不改变校准。这直接削弱了“靠角色人设改善概率”的路径，也说明必须在 LLM 下游加确定性校准层。"),
    ("body", "聚合文献可借鉴：Schoenegger 等（Science Advances 2024）让 12 个 LLM 组成集成，在 31 个二元问题、为期"
             "三个月的预测锦标赛里与 925 名人类对照，集成预测与人类群体统计上不可区分；Halawi 等（NeurIPS 2024）对"
             "同一问题多次提示、用截尾均值聚合，接近甚至部分超越人类竞赛群体——这就是“多采样＋确定性聚合”的成熟"
             "模板。超级预测文献还有条规则：极化聚合只在成员信息相互独立时有效，成员共享信息或经过讨论时反而有害。"),
    ("body", "团队规模的弱信号：两项邻域辩论消融（MADIAVE、PD3）都发现固定预算下智能体从 2 增到 3–5 反而更差，"
             "虽非金融场景，但方向一致：反对扩到七角色，温和支持小委员会。"),

    ("h1", "3　OpenInvest 系统"),
    ("body", "四角色纵向分工（宏观/量化/风控/CIO），单轮、无辩论汇总。核心输出是带概率的路径预测：30/60/90 天涨跌"
             "概率、P10/P90 概率带、路径形状分类，配事后对账闭环。LLM 之外有四类确定性组件：Sanity Check 规则覆盖、"
             "独立快崩防御（按 VIX 分位等哨兵拦买入）、偿付能力兜底、摩擦成本。数据设施是覆盖 2005 年至今的 "
             "walk-forward 重算（1579 个零前视点）＋预注册拟合/样本外验收；重算快照永久保存未校准的原始分布，"
             "校准只在生产出口应用，防止系统吃自己的输出。"),

    ("h1", "4　我的理论框架"),
    ("body", "(1) 三机制分解（图 1）：任何多智能体架构的价值可拆成三种可独立检验的机制——(a) 信息源分解：不同智能体"
             "吃不同数据，这是多智能体相对单智能体的信息论正当性来源；(b) 去相关采样＋聚合：集成降方差，这是受控实验"
             "里增益的主要来源，单模型多采样即可获得，不依赖角色扮演；(c) 确定性流程控制：用规则包裹随机的 LLM 提案，"
             "高风险场景证据最强。角色人设本身不算第四种机制，只有当角色映射到不同信息源时才落入 (a)。"),
    ("fig", ("figure1_mechanisms.png",
             "图 1　多智能体价值的三机制分解：信息源分解(a)、去相关采样＋聚合(b)、确定性流程控制(c)。"
             "角色人设仅当映射到独立信息源时才计入(a)。", 6.0)),
    ("body", "(2) 分层目标：优化层只用 Brier 和覆盖率拟合参数（Sharpe 估计方差太大——十年数据估 SR=0.5、标准误约 "
             "0.34，与估计值同量级；又不可诊断、可被操纵）；决策层用确定性规则把校准概率翻译成仓位（参数极少）；"
             "监控层把损益当输出而非损失函数，看“完整系统减纯确定性策略”的配对差值，共同噪声相消、功效更高。"),
    ("body", "(3) 可验证性（图 2）：收益声明受 t≈SR×√T 约束——夏普 0.5 要约 16 年、夏普 1 也要约 4 年才到 t≈2；"
             "所以任何 MVP 尺度上“系统盈利”和“运气好”在数据上没法区分，而校准声明当下就能验证。主动管理基本定律 "
             "IR≈IC×√breadth 说明可控变量是预测质量（对应校准）和避免负 IC 行为（对应纪律）。按价值可信度排序："
             "防大回撤（跌 50% 要涨 100% 才回本）> 行为纪律 > 成本控制 > 预测 alpha。"),
    ("fig", ("figure2_verifiability.png",
             "图 2　收益类声明在 MVP 时间尺度上为何不可验证（第 4.3 节）。左：达到 t≈2 所需年数随 Sharpe 变化"
             "（t=SR·√T）；右：Sharpe 标准误随样本年数仅以 1/√T 收缩。", 6.3)),

    ("h1", "5　我跑的实验"),
    ("body", "5.1 方向锁消融：开 uptrend 方向锁(A)与生产无锁(B)各跑 912 个决策日，差异只在仓位档位不在方向——A 把"
             "更多决策推向 ACCUMULATE（348 vs 234），HOLD 由 648 降到 506，但 30 天命中率 57.2% vs 55.7%，差值小于"
             "重复跑之间 35–45% 的翻转噪声。结论：这层是“杠杆旋钮”，平移仓位激进度，不改变方向判别的期望正确率。"
             "重心于是转向条件概率表本身的质量。"),
    ("body", "5.2 全历史校准诊断（图 3–5）：整体覆盖率 69–76%，中位偏差约零（说明早先“偏保守”是窗口效应）。但分时代"
             "看，七个时代多数低于 80%；downtrend 整体仅 54%，最差桶 GC=F 2007–2009 downtrend 仅 7%（n=14）——样本"
             "不足时分布呈“虚假精确”（图 3）。图 4 显示条件带宽随有效样本数几乎不变（中位约 9%），但 effective_n=1 时"
             "覆盖率跌到 42%。图 5：条件概率表的 p_below Brier 在所考察七个时代里都不优于同资产无条件分布，这就是要"
             "做“收缩”的动机。"),
    ("fig", ("figure3_era_regime_heatmap.png",
             "图 3　按时代×regime 的 90 天概率带覆盖率（重算自 1579 点）。downtrend 列系统性偏低，最差桶 GC=F "
             "2007–2009 downtrend 仅 7%（n=14），系小样本下的伪精确。", 5.0)),
    ("fig", ("figure4_spurious_precision.png",
             "图 4　“伪精确”（90 天）：条件带宽随有效样本数几乎不变（中位约 9%），但覆盖率在 effective_n=1 时"
             "跌至 42%——模型在稀薄数据上保持同等自信。", 5.4)),
    ("fig", ("figure5_brier_by_era.png",
             "图 5　各时代条件概率表 p_below 的 Brier 与同资产无条件分布对比（90 天）。所考察七个时代条件表均"
             "不更优，构成“向无条件分布收缩”的动机。", 5.7)),
    ("body", "5.3 校准层与样本外验收（图 6、表 2）：两项修正——小样本收缩（把条件分布往同资产无条件收，"
             "λ=eff_n/(eff_n+k)）、带宽扩张（P10/P90 乘系数 γ）。用 2007–2017 拟合、2018–2026 纯样本外验证，三个"
             "窗口覆盖率都进了 [75%, 85%] 且 Brier 全改善，k=80、γ=1.1 写进默认配置。生产端即时效应：GC=F 最差桶 "
             "p_below 由 35% 升到 41%，悲观分位由 −2.3% 修正到 −3.0%，方向和诊断一致。"),
    ("img", ("table2_calibration.png", 6.5)),
    ("fig", ("figure6_coverage_by_window.png",
             "图 6　各窗口概率带覆盖率：未校准 vs 仅带宽扩张 γ=1.1（全历史，n=1579）。表 2 为更严格的样本外验收"
             "口径（叠加小样本收缩 k=80），故样本外数值略高；图 6 用于隔离展示带宽扩张单因子的效应。", 5.5)),
    ("body", "5.4 技术面分析师实验：预注册命中率门——要求每个分析师 30 天方向命中率的 95% 置信区间下界，超过"
             "“同子集基率”和“机械映射基线”的较大者才算 PASS。在 732 条 STANCE 报告上，fundamental/news/sentiment "
             "的 30 天命中率为 50.3%（CI [42.4, 58.3]）/57.4%/64.7%，而方向基率约 71–73%，三者 CI 下界都低于基率，"
             "全部 FAIL，按预注册不进 Phase B。即：单独的分析师人设没能在方向上超过朴素基率。"),
    ("body", "5.5 一次实盘运行（图 7，2026-06-12）：VIX 处近两年 85% 分位，独立防御触发，委员会对两项资产的 "
             "ACCUMULATE 被确定性降级为 HOLD，对黄金的减仓提议被偿付能力规则拦截——确定性层按设计覆盖了 LLM。"
             "同时暴露一个问题：报告把 LLM 自报置信度当主数字，而它几乎全堆在 0.55/0.60/0.65（图 7，n=2100，"
             "中位 0.60），缺校准依据，已列入修订。"),
    ("fig", ("figure7_confidence_distribution.png",
             "图 7　LLM 自报 CONFIDENCE 分布（n=2100 次委员会决策）。取值聚集于 0.05 的整数倍（0.35/0.55/0.60/"
             "0.65），复现了言语化置信度的整数化伪精度。", 5.7)),

    ("h1", "6　架构评估"),
    ("body", "Macro 与 Quant 的划分，只有在两者吃不同信息源时才成立，可由它们概率输出的相关性直接检验（相关大于 "
             "0.9 就该合并）。最弱的环节是“让 LLM 当最终概率聚合者（CIO）”：它既易被输入报告的论断性带偏，自报概率"
             "又系统性过度自信；更好的做法是对各智能体的概率做确定性聚合（均值/中位数/截尾均值），这是当前预期校准"
             "收益最大的单项改动——只换“最终概率怎么产生”，不动 CIO 的执行计划与风险预案。确定性层（风控护栏、快崩"
             "防御、已上线的校准层）是证据支持最强的部分。无辩论设计与证据一致——引入多空辩论是所有可选改动里正当性"
             "最低的，要引入也得过“无害性”门槛。"),
    ("body", "对 RQ1 的回答：不存在被实证确认的最优角色数；七角色和四角色同等地位，弱先验还反对扩张。决策相关的不是"
             "角色计数，而是可分解的独立信息源数量（机制 a）、可聚合的独立概率样本数量（机制 b）。角色数应是信息源"
             "结构的因变量，而非自变量。"),

    ("h1", "7　接下来要做的（见表 3）"),
    ("body", "五阶段预注册实验，都能在现有 walk-forward 上跑：① 委员会折叠对照（等采样预算，四角色＋LLM CIO vs 单"
             "智能体多采样＋确定性聚合）——直接回答 RQ1 的域内版本；② 聚合器消融（用均值/中位数/截尾均值替换 LLM 的"
             "概率裁决）；③ 信息源去相关检验（按宏观/价格因子/流动性/新闻重划角色）；④ 辩论无害性检验（双重非劣才"
             "启用）；⑤ 确定性层强化（把 LLM 概率都当提案，规则可钳制/加宽/否决）。"),
    ("img", ("table3_fivestage.png", 6.5)),
    ("body", "黄金单资产 MVP：只验证管线闭环（概率→仓位→防御→记账四环节贯通），不验证盈利能力（单资产把下注广度"
             "压到 1，盈利更难验证而非更易，这点要先写清防止误读）。选黄金因为它组合内仓位占比最大（49.7%）、概率表"
             "最深、数据现成。重点工作：防御层的黄金语义重检（黄金危机中是双相行为，不能照搬权益语义）、确定性 "
             "sizing、反事实记账、特征审计、三臂归因（买入持有 / 纯确定性策略 / 完整系统，外加 200 日均线作经典基线）。"),

    ("h1", "8　结论"),
    ("body", "没有任何特定角色数是最优的，七角色和四角色同等地位。可检验的设计原则是：按信息源分解智能体、用确定性"
             "方法聚合独立概率样本、用确定性规则层包裹 LLM 提案、避免辩论和人设扭曲概率——这是一套设计纪律，不是固定"
             "蓝图。我系统的无辩论汇总、确定性防御和校准层与证据一致；最该改的是把“LLM 叙事性推导最终置信度”换成确定"
             "性聚合。校准目标和盈利目标用分层框架连接：优化层拟合校准、决策层确定性翻译、监控层配对归因——“经样本外"
             "验证的概率＋强制执行的纪律”，就是这个系统目前能验证的能力边界。"),

    ("h1", "9　局限"),
    ("body", "第 2 节引用的受控研究多在通用推理基准上，不是金融路径概率预测，迁移有缺口（第 7 节实验就是为补这个缺口）；"
             "我同时是系统开发者，有确认偏误风险，预注册是主要缓解；单系统案例外推性有限，可移植的是框架和实验设计，"
             "不是具体参数。再说一遍：我是新手，欢迎指正。"),
    ("refs", "参考（精选）：TradingAgents (arXiv:2412.20138)；Wang et al. ACL 2024；Smit et al. ICML 2024；"
             "Zhang et al. 2025；Tran & Kiela 2026；Choi et al. NeurIPS 2025；Xiong et al. ICLR 2024；"
             "Schoenegger et al. Science Advances 2024；Halawi et al. NeurIPS 2024；FinCon NeurIPS 2024；"
             "Lo 2002；Grinold 1989；Gneiting & Raftery 2007。"),
    ("note", "图表说明：图 3–7 与表 1–3 全部由公开脚本从归档数据 1579 个重算点直接生成，未引入新假设；"
             "图 1–2 为机制示意与解析图。所有图内数字均与正文互校一致。"),
]


def main():
    doc = Document()
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    normal = doc.styles["Normal"]
    normal.font.name = LATIN; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)

    # ---- title block ----
    para(doc, "LLM 多智能体投资系统：到底要几个角色？", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_before=6, space_after=4)
    para(doc, "——我做 OpenInvest 的实验笔记（新手向，欢迎评论区指正）", size=11.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2)
    para(doc, "Sizhuo Long", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2)
    para(doc, "openInvest 项目（github.com/longsizhuo/openInvest）", size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2)
    para(doc, "2026 年 6 月", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=8)

    # ---- intro ----
    for t in INTRO:
        para(doc, t, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.35, indent_first=21, space_after=6)
    # ---- TL;DR ----
    para(doc, "结论速览（太长不看版）", size=12, bold=True, color=DARK, space_before=8, space_after=4)
    for i, (lead, rest) in enumerate(TLDR, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Pt(18); pf.first_line_indent = Pt(-18); pf.space_after = Pt(5); pf.line_spacing = 1.3
        set_run(p.add_run(f"{i}. "), size=10.5, bold=True, color=DARK)
        set_run(p.add_run(lead), size=10.5, bold=True, color=DARK)
        set_run(p.add_run(rest), size=10.5)
    para(doc, "——以下是详细版——", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY,
         space_before=8, space_after=8)

    # ---- body blocks ----
    for kind, payload in BLOCKS:
        if kind == "h1":
            para(doc, payload, size=14, bold=True, color=DARK, space_before=12, space_after=6)
        elif kind == "body":
            para(doc, payload, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.35,
                 indent_first=21, space_after=6)
        elif kind == "fig":
            add_figure(doc, payload[0], payload[1], width=payload[2])
        elif kind == "img":
            add_image(doc, payload[0], width=payload[1])
        elif kind == "refs":
            para(doc, payload, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=GRAY, space_before=8, space_after=4)
        elif kind == "note":
            p = para(doc, payload, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=GRAY, italic=True,
                     space_before=6, space_after=8)
            pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
            for edge in ("top", "bottom"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "4"); b.set(qn("w:color"), "B0B7C3")
                pbdr.append(b)
            sp = pPr.find(qn("w:spacing"))
            (sp.addprevious(pbdr) if sp is not None else pPr.insert(0, pbdr))

    se = doc.settings.element
    z = se.find(qn("w:zoom"))
    if z is not None and z.get(qn("w:percent")) is None:
        z.set(qn("w:percent"), "100")

    doc.save(OUT)
    # report visible char count of body text
    import re
    txt = "".join(p.text for p in doc.paragraphs)
    print("Saved:", OUT)
    print("可见字符(去空白):", len(re.sub(r"\s", "", txt)))
    print("总字符(含空格):", len(txt))


if __name__ == "__main__":
    main()
