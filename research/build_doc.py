#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_doc.py — assemble the enhanced (figures + deep-expansion) version of
"LLM 多智能体投资系统的架构选择与校准优先评估".

Strategy: read the ORIGINAL .docx verbatim (no retyping of Chinese text),
re-emit every paragraph/table with a clean heading hierarchy, embed the 7
recomputed figures at mapped anchors with Chinese captions, and weave in the
data-grounded expansions for sections 5.1, 5.2, 5.4 and 5.5.
"""
import re
import os
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/sessions/busy-optimistic-gates/mnt/uploads/LLM多智能体投资系统的架构选择与校准优先评估.docx"
FIGDIR = "/sessions/busy-optimistic-gates/mnt/WebstormProjects/paper_figures/figures"
OUT = "/sessions/busy-optimistic-gates/mnt/outputs/LLM多智能体投资系统_增订版_配图.docx"

LATIN = "Times New Roman"
CJK = "PingFang SC"        # macOS system CJK font; Word substitutes if absent
GRAY = RGBColor(0x59, 0x5F, 0x66)
DARK = RGBColor(0x20, 0x24, 0x28)

# ---------------------------------------------------------------- font helper
def set_run(run, size=10.5, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = LATIN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN)
    rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:eastAsia"), CJK)
    return run

def para(doc, text="", size=10.5, bold=False, align=None, color=None,
         italic=False, space_after=6, space_before=0, line=None, indent_first=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if line is not None:
        pf.line_spacing = line
    if indent_first is not None:
        pf.first_line_indent = Pt(indent_first)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p

def caption(doc, text):
    para(doc, text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY,
         space_before=3, space_after=10)

def add_figure(doc, fname, cap, width=6.0):
    path = os.path.join(FIGDIR, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    caption(doc, cap)

def add_image(doc, fname, width=6.4, space_before=8, space_after=8):
    """Insert a centered image with NO caption (title is baked into the image)."""
    path = os.path.join(FIGDIR, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run().add_picture(path, width=Inches(width))

def emit_front_matter(doc):
    """Colloquial intro + plain-language TL;DR + reproducibility note (replaces abstract)."""
    for t in INTRO:
        para(doc, t, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.35,
             indent_first=21, space_after=6)
    para(doc, "结论速览（太长不看版）", size=12, bold=True, color=DARK,
         space_before=8, space_after=4)
    for i, (lead, rest) in enumerate(TLDR, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Pt(18); pf.first_line_indent = Pt(-18)
        pf.space_after = Pt(5); pf.line_spacing = 1.3
        set_run(p.add_run(f"{i}. "), size=10.5, bold=True, color=DARK)
        set_run(p.add_run(lead), size=10.5, bold=True, color=DARK)
        set_run(p.add_run(rest), size=10.5)
    # reproducibility note (small, gray, ruled)
    p = para(doc, NOTE_EDITION, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=GRAY,
             italic=True, space_before=8, space_after=10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4"); b.set(qn("w:color"), "B0B7C3")
        pbdr.append(b)
    spacing = pPr.find(qn("w:spacing"))
    (spacing.addprevious(pbdr) if spacing is not None else pPr.insert(0, pbdr))

# ---------------------------------------------------------------- table copy
def copy_table(doc, src_table):
    nrows = len(src_table.rows); ncols = len(src_table.columns)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, row in enumerate(src_table.rows):
        for j, cell in enumerate(row.cells):
            tc = t.cell(i, j)
            tc.text = ""
            txt = cell.text.strip()
            p = tc.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(txt)
            set_run(run, size=8.5, bold=(i == 0), color=(DARK if i == 0 else None))
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "DCE6F1")
                tc._tc.get_or_add_tcPr().append(shd)
    return t

# ============================================================ EXPANSION TEXT
EXP_5_1 = (
    "该消融基于 2026-06-04 的成对重放：在 GC=F 与 NDQ.AX 的周频决策日上，分别运行开启 uptrend "
    "方向锁的 A 链（n=912）与生产用无锁配置的 B 链（n=912）。两链裁决分布的差异集中于仓位档位而非"
    "方向——A 链将更多决策推向 ACCUMULATE（348 比 234，约 +12.5 个百分点），HOLD 由 648 降至 506；"
    "而 30 天命中率在两链间为 57.2% 对 55.7%，其差小于重复 draws 之间 35%–45% 的裁决翻转噪声。"
    "据此该层被定性为“杠杆旋钮”：它平移仓位激进度，却不改变方向判别的期望正确率，对校准目标无可测增量。"
    "这一发现将后续工作重心由规则层转向条件概率表本身的质量。"
)

EXP_5_2_FIGS = (
    "上述诊断可由三组重算图直接读出（均重算自 path_review.jsonl，n=1579，零前视）。图 3 以时代×regime "
    "网格呈现 90 天覆盖率：downtrend 列系统性偏低（整体约 54%），最差桶 GC=F 2007–2009 downtrend 仅 7%"
    "（n=14）。图 4 将该现象归因于样本量——条件带宽随有效样本数几乎不变（中位约 9%），而覆盖率在 "
    "effective_n=1 时跌至 42%，即模型在稀薄数据上保持同等自信，构成伪精确。图 5 比较条件概率表与同资产"
    "无条件分布的 p_below Brier：在所考察的七个时代中，条件表 Brier 均不优于同资产无条件分布（金熊时段"
    "条件表 0.356、无条件 0.333；此处基线取逐点的同资产无条件分布，即收缩目标本身，与本节上文相对全局基率"
    "的 0.228 口径不同），直接为“向同资产无条件分布收缩”这一校准动机提供了证据。"
)

EXP_5_3_FIG = (
    "图 6 以全历史口径直观呈现带宽扩张的方向性效应：仅施加 γ=1.1 即把 30/60/90 天覆盖率由 76%/72%/69% "
    "抬升至 80%/76%/74%。需说明该图为全历史诊断口径，与表 2 报告的拟合（2007–2017）/样本外（2018–2026）"
    "验收口径不同——后者叠加了小样本收缩 k=80，故样本外数值（81%/80%/76%）更高；图 6 仅用于隔离展示带宽扩张"
    "单一因子对覆盖率的影响。"
)

EXP_5_4_TA = (
    "该实验的核心读数为预注册命中率门：对每个分析师，要求其 30 天方向命中率的 95% 置信区间下界超过"
    "“同子集方向基率”与“确定性机械映射基线”二者的较大值，方判 PASS。在 732 条 STANCE 报告（覆盖 GC=F 与 "
    "NDQ.AX、2022 与 2024 两窗）上，三个分析师的 30 天命中率分别为 fundamental 50.3%（CI [42.4%, 58.3%]）、"
    "news 57.4%（CI [49.4%, 65.1%]）、sentiment 64.7%（CI [57.0%, 71.8%]），而对应方向基率约 71%–73%。"
    "三者 CI 下界均显著低于基率，全部判定为 FAIL，据预注册流程不再进入 Phase B。该结果与本文主张方向一致："
    "在本系统的信息结构下，单独的分析师人设未能在方向判别上超越朴素基率，增益（若有）须来自独立样本的聚合"
    "而非人设本身。"
)

EXP_5_5_FIG = (
    "图 7 给出 CONFIDENCE 字段在 2100 次委员会决策上的分布：取值高度聚集于 0.55/0.60/0.65 等 0.05 的整数倍，"
    "中位数 0.60，复现了 Xiong 等 [7] 记录的整数化伪精度。该形态本身即说明此原生数值缺乏校准依据，不宜作为"
    "对外呈现的主数字；以经样本外验收的概率表替代之，已列入修订计划。"
)

SUBTITLE = "——OpenInvest 项目的实验笔记 · 新手向，欢迎评论区指正"

INTRO = [
    "先说这是什么：OpenInvest 是我自己在捣鼓的一个项目——让几个大语言模型（LLM）分别扮演宏观、量化、"
    "风控、CIO 等角色，组成一个“投资委员会”，对资产未来 30/60/90 天的涨跌给出带概率的判断。做的过程中"
    "我一直纠结一个问题：这种多智能体系统，到底要几个角色、怎么分工才最好？是不是角色越多就越聪明？",
    "于是我做了两件事：把 2024–2026 年这方面的论文翻了一遍，再在自己的系统上跑了一批实验来验证。下面是"
    "整理出来的结论。我尽量让每个数字都有真实回测数据撑着（全历史 1579 个零前视重算点，所有图表都能用脚本"
    "复现），但我确实是新手，思路和判断难免有错——如果你看出哪里不对，特别欢迎在评论区告诉我，让我有机会"
    "把这个系统改得更好。",
]

# (lead-in phrase, rest) — rendered as a plain-language numbered TL;DR
TLDR = [
    ("没有“最优角色数”这种东西。",
     "七角色的 TradingAgents 和我的四角色，在证据上其实是平级的——没有任何已发表框架真正做过“角色数量”"
     "的消融实验，角色怎么分基本都是“模仿真实投资公司分工”拍脑袋定的（见表 1）。"),
    ("多智能体的好处，主要来自“多采样＋聚合”，而不是“角色扮演”本身。",
     "在算力对齐的对照实验里，单个模型多跑几次、再用确定性方法把结果聚合起来，就能拿到大部分所谓“多智能体”"
     "的增益（图 1 的机制 b）。"),
    ("让 LLM 自己报概率／置信度，不靠谱。",
     "我系统里这些自报置信度几乎全堆在 0.55／0.60／0.65 这几个数上（图 7），这是已知的“整数化伪精度”，"
     "得在 LLM 外面再加一层确定性校准才能用。"),
    ("辩论不一定有用，还可能更糟。",
     "理论上多轮辩论的期望正确率并不会提升（鞅性质），实际还容易把大家带向“过度自信”，恰好被校准指标狠狠扣分。"),
    ("我给系统加的校准层，是有效的。",
     "用真实数据验证后，30／60／90 天的概率带覆盖率回到了 [75%, 85%] 的目标区间，Brier 也全线改善（图 6、表 2）。"),
    ("能不能赚钱，短期根本测不出来。",
     "统计上，一个夏普 0.5 的策略大约要 16 年数据才能确认（图 2）；当下真正能验证的只有“校准”和“纪律”，不是收益。"),
]

NOTE_EDITION = (
    "图表说明：图 3–图 7 与表 1–表 3 全部由公开脚本（make_figures.py / make_tables.py）直接从归档数据重算生成"
    "（path_review.jsonl，n=1579；委员会决策记录；TA 复现包），未引入任何新的建模假设；图 1–图 2 为三机制"
    "示意与可验证性解析图。所有图内数字均与正文互校一致。"
)

# ---------------- figure anchors (after paragraph index) ----------------
FIG_AFTER_PARA = {
    44: [("figure1_mechanisms.png",
          "图 1　多智能体价值的三机制分解：信息源分解(a)、去相关采样＋聚合(b)、确定性流程控制(c)。"
          "角色人设本身不构成第四种机制，仅当映射到独立信息源时才计入(a)。", 6.0)],
    51: [("figure2_verifiability.png",
          "图 2　收益类声明在 MVP 时间尺度上的统计不可验证性（第 4.3 节）。左：达到 t≈2 所需年数随 Sharpe "
          "变化（t=SR·√T）；右：Sharpe 标准误随样本年数仅以 1/√T 收缩 [Lo 2002]。", 6.3)],
    56: [("figure3_era_regime_heatmap.png",
          "图 3　按时代×regime 的 90 天概率带覆盖率（重算自 path_review.jsonl）。downtrend 列系统性偏低，"
          "最差桶 GC=F 2007–2009 downtrend 仅 7%（n=14），系小样本下的伪精确。", 5.0),
         ("figure4_spurious_precision.png",
          "图 4　“伪精确”现象（90 天窗口）：条件带宽随有效样本数几乎不变（中位约 9%），但覆盖率在 "
          "effective_n=1 时跌至 42%——模型在稀薄数据上保持同等自信。", 5.4),
         ("figure5_brier_by_era.png",
          "图 5　各时代条件概率表 p_below 的 Brier 与同资产无条件分布对比（90 天）。红叉标记条件表劣于无条件"
          "分布的时代；本图所考察七个时代全部如此，构成“向无条件分布收缩”的校准动机。", 5.7)],
    63: [("figure7_confidence_distribution.png",
          "图 7　LLM 自报 CONFIDENCE 字段的分布（n=2100 次委员会决策）。取值聚集于 0.05 的整数倍（0.35/0.55/"
          "0.60/0.65），复现了言语化置信度的整数化伪精度（Xiong 等, ICLR 2024）。", 5.7)],
}
# expansion text paragraphs (after paragraph index) — emitted BEFORE figures of same anchor
TXT_AFTER_PARA = {
    54: [("body", EXP_5_1)],
    56: [("body", EXP_5_2_FIGS)],
    61: [("body", EXP_5_4_TA)],
    63: [("body", EXP_5_5_FIG)],
}
# tables: after this para index, insert table PNG (title baked into image); optional trailing figure
TABLE_AFTER_PARA = {24: 0, 59: 1, 78: 2}
TABLE_PNG = {0: "table1_frameworks.png", 1: "table2_calibration.png", 2: "table3_fivestage.png"}
TABLE_TRAILING = {1: [("figure6_coverage_by_window.png",
                       "图 6　按预测窗口的概率带覆盖率：未校准 vs 仅施加带宽扩张 γ=1.1（全历史，n=1579）。"
                       + "（口径说明见下文；表 2 为更严格的样本外验收口径。）", 5.5)]}

HEADING_RE = re.compile(r"^(\d+)(?:\.(\d+))?[\s　]")

def emit_paragraph(doc, idx, text):
    # Title block
    if idx == 0:
        para(doc, text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=DARK, space_before=6, space_after=4); return
    if idx == 1:  # subtitle -> friendlier
        para(doc, SUBTITLE, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2); return
    if idx in (2, 3, 4):
        para(doc, text, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=2); return
    # Abstract / keywords blocks (5-8) -> replaced by emit_front_matter, skip originals
    if idx in (5, 6, 7, 8):
        return
    # Table caption paragraphs -> skipped (title is baked into the table PNG)
    if idx in TABLE_AFTER_PARA:
        return
    # Headings
    m = HEADING_RE.match(text)
    if text.strip() in ("参考文献",):
        para(doc, text, size=14, bold=True, color=DARK, space_before=14, space_after=8); return
    if m:
        if m.group(2) is None:
            para(doc, text, size=14, bold=True, color=DARK, space_before=14, space_after=8)
        else:
            para(doc, text, size=12, bold=True, color=DARK, space_before=10, space_after=6)
        return
    # References entries (hanging indent)
    if idx >= 93:
        p = para(doc, text, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        return
    # Default body
    para(doc, text, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3,
         indent_first=21, space_after=6)

def main():
    src = Document(SRC)
    paras = [p.text.strip() for p in src.paragraphs]
    doc = Document()
    # base margins
    sec = doc.sections[0]
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    # default style font
    normal = doc.styles["Normal"]
    normal.font.name = LATIN; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)

    for idx, text in enumerate(paras):
        if not text:
            continue
        emit_paragraph(doc, idx, text)
        # inject colloquial intro + TL;DR right after the title block
        if idx == 4:
            emit_front_matter(doc)
        # expansion text after this paragraph
        for kind, t in TXT_AFTER_PARA.get(idx, []):
            if kind == "note":
                p = para(doc, t, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=GRAY,
                         italic=True, space_before=4, space_after=8)
                # subtle top+bottom rule (insert pBdr in schema-correct position: before spacing)
                pPr = p._p.get_or_add_pPr()
                pbdr = OxmlElement("w:pBdr")
                for edge in ("top", "bottom"):
                    b = OxmlElement(f"w:{edge}")
                    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
                    b.set(qn("w:space"), "4"); b.set(qn("w:color"), "B0B7C3")
                    pbdr.append(b)
                spacing = pPr.find(qn("w:spacing"))
                if spacing is not None:
                    spacing.addprevious(pbdr)
                else:
                    pPr.insert(0, pbdr)
            else:
                para(doc, t, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3,
                     indent_first=21, space_after=6)
        # special: 5.3 caliper note after the EXP_5_5? no — add 5.3 note right after fig6 table block
        # tables after this paragraph -> insert table PNG (title baked in)
        if idx in TABLE_AFTER_PARA:
            tidx = TABLE_AFTER_PARA[idx]
            add_image(doc, TABLE_PNG[tidx], width=6.5)
            for fname, cap, w in TABLE_TRAILING.get(tidx, []):
                add_figure(doc, fname, cap, width=w)
                if tidx == 1:
                    para(doc, EXP_5_3_FIG, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line=1.3, indent_first=21, space_after=6)
        # figures after this paragraph
        for fname, cap, w in FIG_AFTER_PARA.get(idx, []):
            add_figure(doc, fname, cap, width=w)

    # fix python-docx default <w:zoom> missing percent (schema completeness)
    se = doc.settings.element
    zoom = se.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    doc.save(OUT)
    print("Saved:", OUT)
    print("paras:", len(paras), "tables:", len(src.tables))

if __name__ == "__main__":
    main()
